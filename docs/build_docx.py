"""Generate a designed Word document (Homies.docx) from the project docs,
including two brand-styled diagrams. Run: python docs/build_docx.py"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ASSETS = os.path.join(HERE, "assets")
os.makedirs(ASSETS, exist_ok=True)

# ---- Brand palette (mirrors theme.dart HomiesColors) ----
ACCENT = "#E85A4F"
ACCENT_STRONG = "#C8473D"
BG = "#F7F6F3"
SURFACE2 = "#FAF9F6"
TEXT = "#2A2730"
TEXT_DIM = "#6B6375"
BORDER = "#E8E5E0"
OK = "#2F855A"
WARN = "#B7791F"
INFO = "#356190"

R_ACCENT = RGBColor(0xE8, 0x5A, 0x4F)
R_ACCENT_STRONG = RGBColor(0xC8, 0x47, 0x3D)
R_TEXT = RGBColor(0x2A, 0x27, 0x30)
R_DIM = RGBColor(0x6B, 0x63, 0x75)
R_WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# =========================================================================
# DIAGRAMS
# =========================================================================
def _box(ax, x, y, w, h, title, sub=None, fill="#FFFFFF", edge=ACCENT,
         tcolor=TEXT, lw=1.6, title_size=11, sub_size=8.5, bold=True):
    ax.add_patch(FancyBboxPatch(
        (x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.6",
        linewidth=lw, edgecolor=edge, facecolor=fill, zorder=2))
    cy = y + h / 2 + (h * 0.13 if sub else 0)
    ax.text(x + w / 2, cy, title, ha="center", va="center",
            fontsize=title_size, color=tcolor,
            fontweight="bold" if bold else "normal", zorder=3)
    if sub:
        ax.text(x + w / 2, y + h / 2 - h * 0.22, sub, ha="center", va="center",
                fontsize=sub_size, color=tcolor, zorder=3)


def _arrow(ax, x1, y1, x2, y2, color=ACCENT_STRONG):
    ax.add_patch(FancyArrowPatch(
        (x1, y1), (x2, y2), arrowstyle="-|>", mutation_scale=14,
        linewidth=1.6, color=color, zorder=1))


def architecture_diagram(path):
    fig, ax = plt.subplots(figsize=(7.4, 9.2))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    ax.text(50, 97, "Architecture at a glance", ha="center",
            fontsize=15, fontweight="bold", color=TEXT)
    ax.text(50, 93, "Flutter mobile app · single source of truth",
            ha="center", fontsize=9.5, color=TEXT_DIM)

    # Main vertical flow (left/centre)
    cx, w = 8, 50
    steps = [
        (82, 8, "main.dart", "App entry · Firebase init", ACCENT),
        (70, 8, "HomiesScope", "InheritedNotifier · exposes state", ACCENT),
        (54, 12, "HomiesState  (ChangeNotifier)",
         "Single source of truth · load() · mutate() · all collections", ACCENT_STRONG),
        (42, 8, "GoRouter", "Public + /app routes", ACCENT),
        (30, 8, "AppShell", "App bar · drawer · bottom nav", ACCENT),
        (16, 9, "Feature screens", "Dashboard · Bills · Cleaning · Chat · …", ACCENT),
    ]
    centers = []
    for (y, h, t, s, e) in steps:
        fill = "#FCEEEC" if e == ACCENT_STRONG else "#FFFFFF"
        _box(ax, cx, y, w, h, t, s, fill=fill, edge=e, title_size=12)
        centers.append((y, h))
    for i in range(len(centers) - 1):
        y_top = centers[i][0]
        y_bot = centers[i + 1][0] + centers[i + 1][1]
        _arrow(ax, cx + w / 2, y_top, cx + w / 2, y_bot)

    # Side panel (right) — services attached to HomiesState
    sx, sw = 66, 30
    state_y = 54 + 12 / 2
    side = [
        (74, "shared_preferences", "JSON persistence", INFO),
        (58, "Firebase Auth", "Email/password · listener", OK),
        (42, "Firestore", "User profiles", OK),
        (26, "seed.dart", "Demo / sample data", WARN),
    ]
    for (y, t, s, e) in side:
        _box(ax, sx, y, sw, 9, t, s, fill="#FFFFFF", edge=e,
             tcolor=TEXT, title_size=10.5, sub_size=8)
        _arrow(ax, cx + w, state_y, sx, y + 4.5, color=BORDER if False else "#B9B3AB")

    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


def feature_map_diagram(path):
    fig, ax = plt.subplots(figsize=(13.6, 8.4))
    fig.patch.set_facecolor(BG)
    ax.set_facecolor(BG)
    ax.set_xlim(0, 100)
    ax.set_ylim(0, 100)
    ax.axis("off")

    ax.text(50, 97, "Feature map", ha="center", fontsize=16,
            fontweight="bold", color=TEXT)
    ax.text(50, 93, "Grouped the way the app's navigation is organised",
            ha="center", fontsize=10, color=TEXT_DIM)

    columns = [
        ("Get started", INFO, [
            "Welcome / landing", "Demo accounts", "Sign up / Sign in",
            "Invites", "Onboarding"]),
        ("Primary", ACCENT, [
            "Dashboard", "Property & lease", "Housemates",
            "Tenant performance *"]),
        ("Money", WARN, [
            "Bills", "Subscriptions", "Groceries", "Necessities"]),
        ("Living together", OK, [
            "Cleaning", "House rules", "Parties", "Messages",
            "House issues", "Complaints", "Find a room  (Marketplace)"]),
        ("Wrap up", TEXT, [
            "Leaving", "End of lease *"]),
    ]

    n = len(columns)
    margin = 3
    gap = 2.5
    cw = (100 - 2 * margin - (n - 1) * gap) / n
    top = 86
    hh = 7        # header height
    fb = 6.4      # feature box height
    fg = 1.6      # gap between feature boxes

    for i, (name, color, feats) in enumerate(columns):
        x = margin + i * (cw + gap)
        # header
        ax.add_patch(FancyBboxPatch(
            (x, top - hh), cw, hh,
            boxstyle="round,pad=0.02,rounding_size=0.6",
            linewidth=0, facecolor=color, zorder=2))
        ax.text(x + cw / 2, top - hh / 2, name, ha="center", va="center",
                fontsize=11, fontweight="bold", color="#FFFFFF", zorder=3)
        # features
        y = top - hh - 3
        for f in feats:
            ax.add_patch(FancyBboxPatch(
                (x, y - fb), cw, fb,
                boxstyle="round,pad=0.02,rounding_size=0.5",
                linewidth=1.3, edgecolor=color, facecolor="#FFFFFF", zorder=2))
            ax.text(x + cw / 2, y - fb / 2, f, ha="center", va="center",
                    fontsize=8.6, color=TEXT, zorder=3, wrap=True)
            y -= (fb + fg)

    ax.text(margin, 4, "*  Leaseholder-only", fontsize=8.5,
            color=TEXT_DIM, style="italic")
    fig.savefig(path, dpi=200, bbox_inches="tight", facecolor=BG)
    plt.close(fig)


# =========================================================================
# WORD HELPERS
# =========================================================================
def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def set_cell_text(cell, text, bold=False, color=None, size=9.5, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def bottom_border(paragraph, color="E85A4F", sz="18"):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def h1(doc, text):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = R_ACCENT_STRONG
    bottom_border(p)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13.5)
    run.font.color.rgb = R_TEXT
    return p


def body(doc, text, dim=False, size=10.5):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = R_DIM if dim else R_TEXT
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = R_TEXT
    return p


def styled_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, htext in enumerate(headers):
        shade(hdr[i], "C8473D")
        set_cell_text(hdr[i], htext, bold=True, color=R_WHITE, size=9.5)
    for r, row in enumerate(rows):
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=9.5)
            shade(cells[i], "FFFFFF" if r % 2 == 0 else "FAF6F5")
    if widths:
        for row in table.rows:
            for i, wdt in enumerate(widths):
                row.cells[i].width = Inches(wdt)
    return table


def chip_line(doc, text):
    """A small uppercase tracked label."""
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(8.5)
    run.font.color.rgb = R_ACCENT_STRONG
    rPr = run._element.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), "30")
    rPr.append(spacing)
    return p


# =========================================================================
# BUILD DOCUMENT
# =========================================================================
def build():
    arch = os.path.join(ASSETS, "architecture.png")
    fmap = os.path.join(ASSETS, "feature_map.png")
    architecture_diagram(arch)
    feature_map_diagram(fmap)

    doc = Document()
    # Page margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.9)
        s.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = R_TEXT

    # ---------------- COVER ----------------
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("🏡  homies")
    run.bold = True
    run.font.size = Pt(42)
    run.font.color.rgb = R_ACCENT_STRONG

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Run a sharehouse without the drama.")
    run.font.size = Pt(16)
    run.font.color.rgb = R_TEXT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Product & technical documentation\nBills · Bond · Chores · Parties · "
        "Marketplace · Move-out — one shared place.")
    run.font.size = Pt(11)
    run.font.color.rgb = R_DIM

    div = doc.add_paragraph()
    div.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bottom_border(div, sz="24")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Flutter mobile app  ·  React web companion")
    run.font.size = Pt(10)
    run.font.color.rgb = R_DIM
    doc.add_page_break()

    # ---------------- CONTENTS ----------------
    h1(doc, "Contents")
    for t in ["1.  Overview", "2.  Architecture", "3.  Feature map",
              "4.  Features — screen by screen", "5.  Data model",
              "6.  Getting started & demo accounts"]:
        bullet(doc, t)
    doc.add_page_break()

    # ---------------- 1. OVERVIEW ----------------
    h1(doc, "1.  Overview")
    body(doc, "Homies is a household-management app for rental sharehouses. It "
              "keeps bills, bond, chores, parties, complaints, the room "
              "marketplace, and the move-in / move-out paperwork in one shared "
              "place — with fair splits and a record everyone can trust.")
    h2(doc, "Two apps, one model")
    styled_table(doc, ["App", "Stack", "Location"], [
        ["Mobile", "Flutter (Dart) + Firebase Auth/Firestore", "homies_mobile/"],
        ["Web", "React + Vite", "src/"],
    ], widths=[1.2, 3.6, 1.7])
    h2(doc, "Roles")
    styled_table(doc, ["Role", "Access"], [
        ["Leaseholder", "Full access + leaseholder-only tools: Tenant "
         "performance and End of lease (termination)."],
        ["Tenant", "Everything except the leaseholder-only tools."],
    ], widths=[1.5, 5.0])
    doc.add_page_break()

    # ---------------- 2. ARCHITECTURE ----------------
    h1(doc, "2.  Architecture")
    body(doc, "All household state lives in one in-memory store that persists "
              "to the device. Firebase is used only for authentication and "
              "profile sync. Every write flows through mutate(), which applies "
              "the change, notifies listeners, and re-persists.")
    doc.add_picture(arch, width=Inches(5.4))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    h2(doc, "Key pieces")
    for b in [
        "HomiesState (ChangeNotifier) — the single source of truth for every "
        "collection plus the session.",
        "HomiesScope (InheritedNotifier) — screens read state via "
        "HomiesScope.of(context) and rebuild on change.",
        "Persistence — JSON in shared_preferences (key homies-mobile-v2); seed "
        "data is the fallback on a fresh install.",
        "Auth — Firebase email/password for real accounts; demo accounts bypass "
        "Firebase via signInAs().",
        "Routing — GoRouter; the authenticated app sits under /app wrapped by "
        "AppShell.",
    ]:
        bullet(doc, b)
    doc.add_page_break()

    # ---------------- 3. FEATURE MAP ----------------
    h1(doc, "3.  Feature map")
    body(doc, "The navigation is organised into Get started, Primary, Money, "
              "Living together, and Wrap up. Leaseholder-only screens are "
              "hidden from tenants.")
    doc.add_picture(fmap, width=Inches(6.6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ---------------- 4. FEATURES ----------------
    h1(doc, "4.  Features — screen by screen")

    sections = [
        ("Onboarding & accounts", [
            ("Welcome / landing — /", "Marketing landing: gradient hero, "
             "feature grid, Marketplace highlight, and entry points to create "
             "an account, sign in, or explore a demo account."),
            ("Demo accounts — /demo", "One-tap, credential-free sign-in as any "
             "pre-seeded housemate. No username or password needed."),
            ("Sign up / Sign in — /signup, /login", "Email + password via "
             "Firebase; new users pick a role."),
            ("Invites — /invite/:code", "Leaseholders invite housemates by "
             "email; invitees accept via a code."),
            ("Onboarding", "Leaseholders configure property + lease; tenants "
             "submit ID/bond/advance proofs and accept house rules."),
        ]),
        ("Primary", [
            ("Dashboard — /app", "At-a-glance home: who's living here, what's "
             "owed, outstanding chores, recent activity."),
            ("Property & lease — /app/property", "Address, type, beds/baths, "
             "features, agent, lease dates, rent, bond + advance weeks."),
            ("Housemates — /app/housemates", "Everyone with role, contact, "
             "move-in date and verification status."),
            ("Tenant performance — /app/performance  (leaseholder-only)",
             "Per-tenant scorecard: chore rate, overdue/excused, bills paid vs "
             "owed, complaints, parties hosted, overall standing."),
        ]),
        ("Money", [
            ("Bills — /app/bills", "One-off and recurring bills with category, "
             "amount, due date, equal/custom split, who-paid tracking, proof, "
             "and schedules."),
            ("Subscriptions — /app/subscriptions", "Shared recurring services: "
             "amount, cadence, payer, participants, split."),
            ("Groceries — /app/groceries", "Shared runs: total, payer, split, "
             "date, optional receipt."),
            ("Necessities — /app/necessities", "Small shared items: item, "
             "payer, amount, date — shared or personal."),
        ]),
        ("Living together", [
            ("Cleaning — /app/cleaning", "Roster (day/area/assignee) + tasks "
             "with due dates, photo proof, or a logged excuse."),
            ("House rules — /app/rules", "Shared rules with author + timestamp; "
             "accepted during onboarding."),
            ("Parties — /app/parties", "Plan events with date, time, host, "
             "notes; housemates RSVP."),
            ("Messages — /app/messages", "House group chat + DMs, with text "
             "posts and polls (single/multi, live tallies)."),
            ("House issues — /app/issues", "Maintenance log with category, "
             "description, optional photo, open/fixed status."),
            ("Complaints — /app/complaints", "Formal complaints: against whom, "
             "reason, severity, open/resolved."),
            ("Find a room (Marketplace) — /app/listings", "Two-sided "
             "marketplace: list a room or post what you're after, browse "
             "listings, and share only chosen details on a match."),
        ]),
        ("Wrap up (moving out)", [
            ("Leaving — /app/leaving", "Give notice: dates, reason, bond "
             "return, itemized deductions the tenant can agree to."),
            ("End of lease — /app/termination  (leaseholder-only)", "Plan "
             "termination: itemized expenses, split mode, notes."),
        ]),
    ]
    for group, items in sections:
        chip_line(doc, group)
        for title, desc in items:
            p = doc.add_paragraph()
            r = p.add_run(title)
            r.bold = True
            r.font.size = Pt(11)
            r.font.color.rgb = R_TEXT
            p2 = doc.add_paragraph()
            r2 = p2.add_run(desc)
            r2.font.size = Pt(10)
            r2.font.color.rgb = R_DIM
            p2.paragraph_format.space_after = Pt(6)
    doc.add_page_break()

    # ---------------- 5. DATA MODEL ----------------
    h1(doc, "5.  Data model")
    body(doc, "Entities are plain classes with toJson / fromJson "
              "(state/models.dart). The web app mirrors the same shapes.")

    h2(doc, "Core")
    styled_table(doc, ["Entity", "Key fields"], [
        ["User", "id, name, role (leaseholder|tenant), email, phone, "
         "moveIn/Out, bondPaid/Amount, advanceRentPaid, docVerified, "
         "acceptedRulesAt, pending, submissions"],
        ["Property", "address, type, beds, baths, features, agent, lease "
         "dates, rentAmount/cadence, bondWeeks, advanceWeeks, setupComplete"],
        ["Session", "userId, pendingSignup"],
    ], widths=[1.3, 5.2])

    h2(doc, "Money")
    styled_table(doc, ["Entity", "Key fields"], [
        ["Bill", "title, category, amount, dueDate, issuedBy, split, shares, "
         "status, paidBy, scheduleId?, proof?"],
        ["BillSchedule", "cadence, customDays, cycleStart, nextDueDate, "
         "estimatedAmount, splitMethod, participants, active"],
        ["Subscription", "name, amount, cadence, payer, participants, split, "
         "shares"],
        ["Grocery", "title, total, payer, mode, split, shares, date, receipt?"],
        ["Necessity", "item, mode, payer, amount, date"],
    ], widths=[1.3, 5.2])

    h2(doc, "Living together")
    styled_table(doc, ["Entity", "Key fields"], [
        ["HouseRule", "text, addedBy, addedAt"],
        ["CleaningRosterEntry", "day, area, assignee"],
        ["CleaningTask", "task, assignee, dueDate, done, photo?, excuse?, "
         "completedAt?"],
        ["Party", "title, date, time, host, notes, responses, status"],
        ["Message / MessagePoll", "from, text, at, type; poll: question, "
         "multi, closed, options, votes"],
        ["Complaint", "against, from, reason, severity, date, status"],
        ["Issue", "title, category, description, photo?, raisedBy, status, "
         "fixedAt/By"],
    ], widths=[1.7, 4.8])

    h2(doc, "Marketplace & moving out")
    styled_table(doc, ["Entity", "Key fields"], [
        ["Listing", "type (tenant-wanted|room-wanted), by, title, suburb, "
         "rent|budget, availableFrom, description, status"],
        ["ListingInterest", "listingId, from, to, message, sharedFields, "
         "status (pending|accepted|declined)"],
        ["Notice", "userId, givenAt, leaveDate, reason, bondReturn, "
         "deductions[], deductionExplanation, tenantAgreed"],
        ["TerminationPlan", "expenses[], splitMode, customShares, notes"],
    ], widths=[1.5, 5.0])
    doc.add_page_break()

    # ---------------- 6. GETTING STARTED ----------------
    h1(doc, "6.  Getting started & demo accounts")
    h2(doc, "Run the mobile app")
    for b in ["cd homies_mobile", "flutter pub get", "flutter run"]:
        p = doc.add_paragraph()
        r = p.add_run("   " + b)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = R_TEXT

    h2(doc, "Explore instantly with demo accounts")
    for b in [
        "Launch the app → Welcome screen.",
        "Tap “Just looking? Explore with a demo account” (or open /demo).",
        "Pick any seeded housemate — leaseholder or tenant — to sign in with no "
        "username or password.",
    ]:
        bullet(doc, b)

    h2(doc, "Seed / demo data")
    for b in [
        "1 property — 12 Marrickville Road, Marrickville NSW 2204 (4-bed).",
        "2 leaseholders — Maya Chen, Daniel Okafor.",
        "3 tenants — Priya Sharma, Tom Becker, Aisha Rahman (onboarding).",
        "A handful of group-chat posts.",
    ]:
        bullet(doc, b)
    body(doc, "Note: persisted state overrides the seed; clear app storage or "
              "call HomiesState.reset() to restore it.", dim=True, size=9.5)

    h2(doc, "Run the web app")
    for b in ["npm install", "npm run dev"]:
        p = doc.add_paragraph()
        r = p.add_run("   " + b)
        r.font.name = "Consolas"
        r.font.size = Pt(10)
        r.font.color.rgb = R_TEXT

    out = os.path.join(HERE, "Homies.docx")
    doc.save(out)
    print("Saved:", out)
    print("Diagrams:", arch, fmap)


if __name__ == "__main__":
    build()
